import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
from dataclasses import dataclass
from typing import List
import unicodedata
import re


@dataclass
class PageText:
    page_number: int
    text: str


class DocumentLoader:
    """
    โหลดเอกสาร (PDF หรือ Word)
    - รองรับทั้ง PDF (.pdf) และ Word (.docx)
    - ตรวจจับประเภทไฟล์อัตโนมัติจาก bytes
    - คืนค่ารายการ PageText(page_number, text)
    """

    # ===========================
    # 🔥 ฟังก์ชัน CLEAN (ใหม่)
    # ===========================
    def _clean_text(self, text: str) -> str:
        if not text:
            return ""

        # 1) Normalize Unicode (ป้องกันอักขระเพี้ยนแบบ การ → กำร)
        text = unicodedata.normalize("NFC", text)

        # 2) ลบ zero-width / control chars
        text = re.sub(r"[\u200b\u200c\u200d\u2060]", "", text)

        # 3) ลบอักขระ control ที่ไม่ใช่ newline
        text = "".join(ch for ch in text if ch == "\n" or not unicodedata.category(ch).startswith("C"))

        # 4) รวมช่องว่างซ้ำ
        text = re.sub(r"[ \t]+", " ", text)

        return text.strip()

    # ===========================
    # PUBLIC API (เหมือนเดิม)
    # ===========================
    def load_from_bytes(self, file_bytes: bytes) -> List[PageText]:
        if file_bytes[:4] == b"%PDF":
            return self._load_pdf(file_bytes)
        elif file_bytes[:2] == b"PK":
            return self._load_docx(file_bytes)
        else:
            raise RuntimeError("ไม่สามารถระบุประเภทไฟล์ได้ (ไม่ใช่ PDF หรือ DOCX)")

    # -------------------------------
    # 🔹 PDF (ใส่ CLEAN เข้าไปแล้ว)
    # -------------------------------
    def _load_pdf(self, pdf_bytes: bytes) -> List[PageText]:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        except Exception as e:
            raise RuntimeError(f"ไม่สามารถเปิดไฟล์ PDF ได้: {e}")

        pages: List[PageText] = []

        for i in range(len(doc)):
            page = doc.load_page(i)
            blocks = page.get_text("blocks")
            lines: List[str] = []

            for b in blocks:
                text = b[4]
                if not text:
                    continue

                for line in text.split("\n"):
                    clean = self._clean_text(line)   # ✅ CLEAN ตรงนี้
                    if clean:
                        lines.append(clean)

            pages.append(
                PageText(
                    page_number=i + 1,
                    text="\n".join(lines)
                )
            )

        doc.close()
        return pages

    # -------------------------------
    # 🔹 Word (ใส่ CLEAN เข้าไปแล้ว)
    # -------------------------------
    def _load_docx(self, docx_bytes: bytes) -> List[PageText]:
        try:
            doc = Document(BytesIO(docx_bytes))
        except Exception as e:
            raise RuntimeError(f"ไม่สามารถเปิดไฟล์ Word ได้: {e}")

        pages: List[PageText] = []
        current_lines: List[str] = []
        page_number = 1

        body = doc.element.body
        para_iter = iter(doc.paragraphs)

        for element in body:
            tag = element.tag.split("}")[-1]

            # =========================
            # 1️⃣ Paragraph
            # =========================
            if tag == "p":
                try:
                    p = next(para_iter)
                except StopIteration:
                    continue

                text = self._clean_text(" ".join(p.text.split()))  # ✅ CLEAN ตรงนี้

                if text:
                    current_lines.append(text)

                # page break
                has_page_break = any(
                    r.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                    ) == "page"
                    for r in element.xpath(".//w:br")
                )

                if has_page_break:
                    pages.append(
                        PageText(
                            page_number=page_number,
                            text="\n".join(current_lines).strip(),
                        )
                    )
                    page_number += 1
                    current_lines = []

            # =========================
            # 2️⃣ Table
            # =========================
            elif tag == "tbl":
                table = None
                for t in doc.tables:
                    if t._tbl == element:
                        table = t
                        break

                if table:
                    current_lines.append("[TABLE]")

                    for row in table.rows:
                        cells = []
                        for cell in row.cells:
                            cell_text = " ".join(
                                self._clean_text(p.text.strip())
                                for p in cell.paragraphs
                                if p.text.strip()
                            )
                            cells.append(cell_text)

                        current_lines.append(" | ".join(cells))

                    current_lines.append("[/TABLE]")

        # หน้าสุดท้าย
        if current_lines:
            pages.append(
                PageText(
                    page_number=page_number,
                    text="\n".join(current_lines).strip(),
                )
            )

        return pages
